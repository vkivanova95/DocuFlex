import os
from datetime import date
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT


def generate_annex_standard(data, request_obj, extra_conditions=None):
    doc = init_document()
    add_annex_header(doc, data, request_obj)
    add_parties_info(doc, request_obj.client)
    add_intro_paragraph(doc, data, request_obj)
    add_change_sections(doc, data, request_obj, extra_conditions)
    add_closing_paragraphs(doc, data, request_obj)
    add_signatures(doc, request_obj.client, request_obj.maker.get_full_name())
    return save_document(doc, request_obj.request_number)


def init_document():
    doc = Document()
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Times New Roman'
    font.size = Pt(12)
    style.paragraph_format.space_after = Pt(6)
    return doc


def add_annex_header(doc, data, request_obj):
    heading = doc.add_paragraph(f'Анекс № {data["annex_number"]}')
    heading.runs[0].bold = True
    heading.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    subheading = doc.add_paragraph(f'към Договор за кредит № {request_obj.loan_agreement.contract_number}')
    subheading.runs[0].bold = True
    subheading.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    doc.add_paragraph(f'Днес, {data["annex_date"].strftime("%d.%m.%Y")} в гр. {data["city"]}, между:')


def add_parties_info(doc, client):
    p = doc.add_paragraph()
    p.add_run("БАНКА Х").bold = True
    p.add_run(", вписано в ТРРЮЛНЦ, ЕИК 111222333, със седалище и адрес на управление: гр. София, в качеството "
              "й на Кредитор, наричано по-долу за краткост “Банката” и/или „Кредитора“, от една страна, и")
    address = f"гр. {client.town}"
    if client.district:
        address += f", кв. {client.district}"
    if client.street:
        address += f", ул. {client.street}"
    if client.number:
        address += f", № {client.number}"
    if client.block:
        address += f", бл. {client.block}"
    if client.floor:
        address += f", ет. {client.floor}"
    if client.apartment:
        address += f", ап. {client.apartment}"

    representatives = client.representative1
    if client.representative2:
        representatives += f" и {client.representative2}"

    p = doc.add_paragraph()
    p.add_run(client.name).bold = True
    p.add_run(f", вписано в ТРРЮЛНЦ, ЕИК {client.eik}, със седалище и адрес на управление: {address}, представлявано от {representatives}")
    p.paragraph_format.space_after = Pt(6)


def add_intro_paragraph(doc, data, request_obj):
    doc.add_paragraph('всички съвместно наричани по-долу “Страните”')
    doc.add_paragraph(f'на основание постигнато споразумение се сключи настоящия Анекс № {data["annex_number"]}, '
                      f'с който Страните правят следните изменения и допълнения в сключения между тях Договор за '
                      f'кредит № {request_obj.loan_agreement.contract_number} и всички последващи анекси към него '
                      f'(Договора за кредит и анексите/ допълнителните споразумения към него заедно наричани по-долу '
                      f'за краткост „Договор/а за кредит” и/или „Договор/а”), а именно:')


def add_change_sections(doc, data, request_obj, extra_conditions):
    section_I, section_II, section_III, section_IV, section_V = [], [], [], [], []

    # SECTION I – Промени в размера и валутата на кредита
    counter_I = 1
    if data['new_amount_reduction']:
        section_I.append(f"{counter_I}. Ограничава се ползването на кредита, като максималния размер става "
                         f"{data['new_amount_reduction']}.")
        counter_I += 1
    if data['amount_increase'] and data['new_amount_increase']:
        section_I.append(f"{counter_I}. Увеличава се размера на кредита с {data['amount_increase']}  "
                         f"(„Сума на увеличението“), като максималния разрешен размер на кредита става {data['new_amount_increase']}.")
        counter_I += 1

    # SECTION II – Промени в условията за усвояване
    counter_II = 1
    if data['stop_disbursement']:
        section_II.append(
            f"{counter_II}. Прекратява се усвояването по кредита и задължението на Банката за предоставяне на средства от кредита отпада.")
        counter_II += 1

    # SECTION III – Срокове и погасяване
    counter_III = 1
    if data['new_disbursement_date']:
        section_III.append(
            f"{counter_III}.Удължаване на срока на усвояване – нова дата: {data['new_disbursement_date'].strftime('%d.%m.%Y')} г.")
        counter_III += 1
    if data['new_repayment_date']:
        section_III.append(
            f"{counter_III}.Удължаване на срока на издължаване – нова дата: {data['new_repayment_date'].strftime('%d.%m.%Y')} г.")
        counter_III += 1
    if data['effective_from']:
        section_III.append(
            f"{counter_III}. Във връзка с договорените с настоящия Анекс промени в сроковете с действие считано от "
            f"{data['effective_from'].strftime('%d.%m.%Y')} г., Страните се съгласяват, че за периода от "
            f"{data['effective_from'].strftime('%d.%m.%Y')} г. до отразяване на промяната по този Анекс в системата "
            f"на Банката се дължи уговорената редовна лихва по кредита.")
        counter_III += 1
    if data['repayment_plan']:
        section_III.append(f"{counter_III}. Кредитът се издължава при спазване на погасителен план за намаляване на "
                           f"главницата по кредита, съгласно Приложение, неразделна част от настоящия Анекс.")
        counter_III += 1
    if data['new_ceiling']:
        section_III.append(f"{counter_III}. Кредитополучателят се задължава да спазва следната схема за ограничения "
                           f"в ползването и временно  намаление на максималния разрешен размер на кредита, съгласно "
                           f"Приложение, неразделна част от настоящия Анекс.")
        counter_III += 1

    # SECTION IV – Лихви и такси
    counter_IV = 1
    if data['new_interest']:
        section_IV.append(f"{counter_IV}. Редовната лихва по кредита става {data['new_interest']}")
        counter_IV += 1
    if data['fee_review']:
        section_IV.append(f"{counter_IV}. На датата на подписване на този Анекс Кредитополучателят дължи такса "
                          f"разглеждане в размер на {data['fee_review']}.")
        counter_IV += 1
    if data['fee_management']:
        section_IV.append(f"{counter_IV}. Такса за управление и обработка (обслужване) на кредита става"
                          f" {data['fee_commitment']} годишно, платима в началото на всяко календарно тримесечие, считано от датата на "
                          f"подписване на настоящия анекс.")
        counter_IV += 1
    if data['fee_commitment']:
        section_IV.append(
            f"{counter_IV}. Комисиона за ангажимент става {data['fee_commitment']} годишно, изчислена върху "
            f"размера на разрешената за съответния период, но неизползвана част от кредита, "
            f"като начисляването започва от деня на подписване на този Анекс, платима месечно, "
            f"в последния ден от съответния месец.")
        counter_IV += 1
    if data['no_fees']:
        section_IV.append(f"{counter_IV}. Кредитополучателят не дължи такси по кредита.")
        counter_IV += 1

    # SECTION V – Други условия
    counter_V = 1
    if data['other_V']:
        section_V.append(f"{counter_V}. {data['other_V']}")
        counter_V += 1
    if extra_conditions:
        for cond in extra_conditions:
            section_V.append(f"{counter_V}. {cond}")
            counter_V += 1


    def add_section(title, *contents):
        doc.add_paragraph()
        section_texts = [text for text in contents if text]

        if not section_texts:
            # Добавяме "Няма" на същия ред като заглавието
            p = doc.add_paragraph()
            run_title = p.add_run(title + " ")
            run_title.bold = True
            run_none = p.add_run("Няма")
            p.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
            return

        # Стандартен случай с отделен параграф за заглавието
        title_paragraph = doc.add_paragraph()
        title_run = title_paragraph.add_run(title)
        title_run.bold = True
        title_paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
        if section_texts:
            for line in section_texts:
                p = doc.add_paragraph(line)
        else:
            p = doc.add_paragraph("Няма")

    add_section('I. Промени в размера и валутата на кредита:', *section_I)
    add_section('II. Промени в условията за усвояване:', *section_II)
    add_section('III. Промени в срокове и начин на погасяване:', *section_III)
    add_section('IV. Промени в лихви и такси:', *section_IV)
    add_section('V. Други условия:', *section_V)


def add_closing_paragraphs(doc, data, request_obj):
    doc.add_paragraph()
    doc.add_paragraph(
        f'Всички клаузи от Договор за кредит № {request_obj.loan_agreement.contract_number} и анексите към него, които не са променени с настоящия Анекс, остават непроменени и запазват действието си между Страните.'),
    doc.add_paragraph(
        f'Настоящият Анекс № {data["annex_number"]} се подписа в два еднообразни екземпляра - по един за всяка от Страните и е неразделна част към Договор за кредит № {request_obj.loan_agreement.contract_number}'),

    doc.add_paragraph()
    doc.add_paragraph(
        f'Всички клаузи от Договор за кредит № {request_obj.loan_agreement.contract_number} и анексите към него, които не са променени с настоящия Анекс, остават непроменени и запазват действието си между Страните.')
    doc.add_paragraph(
        f'Настоящият Анекс № {data["annex_number"]} се подписа в два еднообразни екземпляра - по един за всяка от Страните и е неразделна част към Договор за кредит № {request_obj.loan_agreement.contract_number}')


def add_signatures(doc, client, maker_name):
    doc.add_paragraph()
    doc.add_paragraph()
    table = doc.add_table(rows=4, cols=2)
    table.autofit = False
    widths = [Inches(3.25), Inches(3.25)]
    for row in table.rows:
        for idx, cell in enumerate(row.cells):
            cell.width = widths[idx]

    table.cell(0, 0).paragraphs[0].add_run('За БАНКАТА:').bold = True
    table.cell(0, 1).paragraphs[0].add_run('За КРЕДИТОПОЛУЧАТЕЛЯ:').bold = True

    table.cell(1, 0).paragraphs[0].add_run('БАНКА Х')
    table.cell(1, 1).paragraphs[0].add_run(client.name)

    table.cell(2, 0).paragraphs[0].add_run('1. ..............................................')
    table.cell(2, 0).add_paragraph('Търговски пълномощник')
    table.cell(2, 1).paragraphs[0].add_run('1. ................................................')
    table.cell(2, 1).add_paragraph(f'{client.representative1}')
    doc.add_paragraph()
    table.cell(3, 0).paragraphs[0].add_run('2. ..............................................')
    table.cell(3, 0).add_paragraph('Търговски пълномощник')
    if client.representative2:
        table.cell(3, 1).paragraphs[0].add_run('2. ................................................')
        table.cell(3, 1).add_paragraph(client.representative2)
    else:
        table.cell(3, 1).paragraphs[0].add_run('')
        table.cell(3, 1).add_paragraph('')

    doc.add_paragraph()
    p = doc.add_paragraph()
    run = p.add_run('Изготвил: ')
    run.italic = True
    p.add_run(maker_name)

def save_document(doc, request_number):
    path = f'media/annexes/annex_{request_number}_{date.today()}.docx'
    os.makedirs(os.path.dirname(path), exist_ok=True)
    doc.save(path)
    return path


def generate_annex_deletion(data, request_obj, extra_conditions=None):
    doc = init_document()
    add_annex_header(doc, data, request_obj)
    add_parties_info(doc, request_obj.client)

    doc.add_paragraph('всички съвместно наричани по-долу “Страните”')
    doc.add_paragraph(f'на основание постигнато споразумение се сключи настоящия Анекс № {data["annex_number"]}, '
                      f'с който Страните правят следните изменения и допълнения в сключения между тях Договор за '
                      f'кредит № {request_obj.loan_agreement.contract_number} и всички последващи анекси към него '
                      f'(Договора за кредит и анексите/ допълнителните споразумения към него заедно наричани по-долу '
                      f'за краткост „Договор/а за кредит” и/или „Договор/а”), а именно:')

    doc.add_paragraph()
    doc.add_paragraph(f'    1. Във връзка с извършено частично предсрочно погасяване по главница в размер на '
                      f'{data["repaid_amount"]} и съгласно условията на Договора за кредит, след заплащане на '
                      f'всички дължими нотариални и държавни такси за извършване на заличаването (изчислени върху '
                      f'размера на частично погасената сума посочен по-горе), Банката ще даде в предвидената от '
                      f'закона форма съгласие за частично заличаване на учредената в нейна полза договорна ипотека, '
                      f'обективирана в нотариален акт №{data["deed_number"]} само върху следното обезпечение:')
    doc.add_paragraph(f'    - {data["collateral_description"]}')

    doc.add_paragraph()
    doc.add_paragraph(f'Всички клаузи от Договор за кредит № {request_obj.loan_agreement.contract_number} и анексите към него, които не са променени с настоящия Анекс, остават непроменени и запазват действието си между Страните.')
    doc.add_paragraph(f'Настоящият Анекс № {data["annex_number"]} се подписа в два еднообразни екземпляра - по един за всяка от Страните и е неразделна част към Договор за кредит № {request_obj.loan_agreement.contract_number}')

    add_signatures(doc, request_obj.client, request_obj.maker.get_full_name())
    return save_document(doc, request_obj.request_number)
